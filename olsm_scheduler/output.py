"""Generate the 7 output files for the OLSM schedule build."""

from __future__ import annotations

from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side

from .models import (
    ALL_PERIODS, Gender, Period, ScheduleData, SlotType,
)


HEADER_FONT = Font(bold=True, size=11)
HEADER_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
HEADER_FONT_WHITE = Font(bold=True, size=11, color="FFFFFF")
THIN_BORDER = Border(
    left=Side(style="thin"), right=Side(style="thin"),
    top=Side(style="thin"), bottom=Side(style="thin"),
)
ROTATION_FILL = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
SEM_FILL = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
LUNCH_FILL = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")


def _style_header(ws, row: int, max_col: int):
    for col in range(1, max_col + 1):
        cell = ws.cell(row=row, column=col)
        cell.font = HEADER_FONT_WHITE
        cell.fill = HEADER_FILL
        cell.border = THIN_BORDER
        cell.alignment = Alignment(horizontal="center", wrap_text=True)


def _auto_width(ws):
    for col in ws.columns:
        max_len = 0
        col_letter = col[0].column_letter
        for cell in col:
            if cell.value:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_letter].width = min(max_len + 4, 40)


def generate_master_grid(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Master Schedule Grid"

    headers = ["Teacher", "Department"] + [p.value for p in ALL_PERIODS]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header(ws, 1, len(headers))

    row = 2
    for t in sorted(data.teachers.values(), key=lambda t: (t.department, t.name)):
        ws.cell(row=row, column=1, value=t.name)
        ws.cell(row=row, column=2, value=t.department)
        for i, p in enumerate(ALL_PERIODS):
            cell = ws.cell(row=row, column=3 + i)
            slot = t.schedule.get(p)
            class_name = t.assigned_classes.get(p, "")
            if slot == SlotType.ROTATION:
                cell.value = class_name or "ROTATION"
                cell.fill = ROTATION_FILL
            elif slot == SlotType.SEM:
                cell.value = "SEM"
                cell.fill = SEM_FILL
            elif slot == SlotType.LUNCH:
                cell.value = "LUNCH"
                cell.fill = LUNCH_FILL
            elif slot == SlotType.PREP:
                cell.value = "PREP"
            elif slot == SlotType.TEACHING:
                cell.value = class_name
            else:
                cell.value = ""
            cell.border = THIN_BORDER
        row += 1

    _auto_width(ws)
    path = output_dir / "MASTER_SCHEDULE_GRID.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_student_schedules(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Student Schedules"

    headers = ["Student ID", "Last Name", "First Name", "Grade", "Gender",
               "Rotation Sec", "Rotation Q"] + [p.value for p in ALL_PERIODS] + ["Status"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header(ws, 1, len(headers))

    row = 2
    for s in sorted(data.students.values(), key=lambda s: (s.grade, s.gender.value, s.last_name)):
        ws.cell(row=row, column=1, value=s.student_id)
        ws.cell(row=row, column=2, value=s.last_name)
        ws.cell(row=row, column=3, value=s.first_name)
        ws.cell(row=row, column=4, value=s.grade)
        ws.cell(row=row, column=5, value=s.gender.value)
        ws.cell(row=row, column=6, value=s.rotation_section or "")
        ws.cell(row=row, column=7, value=s.rotation_subcohort or "")

        for i, p in enumerate(ALL_PERIODS):
            cell = ws.cell(row=row, column=8 + i)
            val = s.schedule.get(p, "")
            cell.value = val
            if val == "SEM":
                cell.fill = SEM_FILL
            elif val == "LUNCH":
                cell.fill = LUNCH_FILL
            elif "Rotation" in str(val):
                cell.fill = ROTATION_FILL
            cell.border = THIN_BORDER

        status = "VALID" if s.is_schedule_complete() and not s.conflicts else "INCOMPLETE"
        ws.cell(row=row, column=8 + len(ALL_PERIODS), value=status)
        row += 1

    _auto_width(ws)
    path = output_dir / "STUDENT_SCHEDULES.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_class_rosters(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    for sec in sorted(data.sections.values(), key=lambda s: s.section_id):
        if not sec.enrolled_students:
            continue
        title = sec.section_id[:31]
        ws = wb.create_sheet(title=title)

        ws.cell(row=1, column=1, value="Section:")
        ws.cell(row=1, column=2, value=sec.section_id)
        ws.cell(row=2, column=1, value="Course:")
        ws.cell(row=2, column=2, value=sec.course_code)
        ws.cell(row=3, column=1, value="Teacher:")
        ws.cell(row=3, column=2, value=sec.teacher_name)
        ws.cell(row=4, column=1, value="Period:")
        ws.cell(row=4, column=2, value=sec.period.value if sec.period else "N/A")
        ws.cell(row=5, column=1, value="Enrollment:")
        ws.cell(row=5, column=2, value=f"{sec.enrollment}/{sec.capacity}")

        headers = ["#", "Student ID", "Last Name", "First Name", "Grade", "Gender"]
        for col, h in enumerate(headers, 1):
            ws.cell(row=7, column=col, value=h)
        _style_header(ws, 7, len(headers))

        row = 8
        for i, sid in enumerate(sorted(sec.enrolled_students), 1):
            student = data.students.get(sid)
            ws.cell(row=row, column=1, value=i)
            ws.cell(row=row, column=2, value=sid)
            if student:
                ws.cell(row=row, column=3, value=student.last_name)
                ws.cell(row=row, column=4, value=student.first_name)
                ws.cell(row=row, column=5, value=student.grade)
                ws.cell(row=row, column=6, value=student.gender.value)
            row += 1

        _auto_width(ws)

    if not wb.sheetnames:
        wb.create_sheet("Empty")

    path = output_dir / "CLASS_ROSTERS.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_teacher_schedules(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    for t in sorted(data.teachers.values(), key=lambda t: t.name):
        title = t.name[:31]
        ws = wb.create_sheet(title=title)

        ws.cell(row=1, column=1, value="Teacher:")
        ws.cell(row=1, column=2, value=t.name)
        ws.cell(row=2, column=1, value="Department:")
        ws.cell(row=2, column=2, value=t.department)
        ws.cell(row=3, column=1, value="Teaching Periods:")
        ws.cell(row=3, column=2, value=t.teaching_period_count())
        ws.cell(row=4, column=1, value="Overload:")
        ws.cell(row=4, column=2, value="YES" if t.overload_approved else "No")

        headers = ["Period", "Assignment", "Type"]
        for col, h in enumerate(headers, 1):
            ws.cell(row=6, column=col, value=h)
        _style_header(ws, 6, len(headers))

        row = 7
        for p in ALL_PERIODS:
            ws.cell(row=row, column=1, value=p.value)
            slot = t.schedule.get(p)
            class_name = t.assigned_classes.get(p, "")
            cell_assign = ws.cell(row=row, column=2, value=class_name)
            cell_type = ws.cell(row=row, column=3, value=slot.value if slot else "FREE")

            if slot == SlotType.ROTATION:
                cell_assign.fill = ROTATION_FILL
            elif slot == SlotType.SEM:
                cell_assign.fill = SEM_FILL
            elif slot == SlotType.LUNCH:
                cell_assign.fill = LUNCH_FILL
            row += 1

        sections = data.teacher_sections(t.name)
        if sections:
            row += 1
            ws.cell(row=row, column=1, value="Class Rosters:")
            row += 1
            for sec in sections:
                ws.cell(row=row, column=1, value=sec.section_id)
                ws.cell(row=row, column=2, value=f"{sec.enrollment}/{sec.capacity} students")
                ws.cell(row=row, column=3, value=sec.period.value if sec.period else "N/A")
                row += 1

        _auto_width(ws)

    path = output_dir / "TEACHER_SCHEDULES.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_audit_log(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Audit Log"

    headers = ["Timestamp", "Step", "Check ID", "Check Name", "Result", "Detail"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header(ws, 1, len(headers))

    fail_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
    warn_fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
    pass_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")

    for row, entry in enumerate(data.audit_log, 2):
        ws.cell(row=row, column=1, value=entry.timestamp)
        ws.cell(row=row, column=2, value=entry.step)
        ws.cell(row=row, column=3, value=entry.check_id)
        ws.cell(row=row, column=4, value=entry.check_name)
        result_cell = ws.cell(row=row, column=5, value=entry.result)
        ws.cell(row=row, column=6, value=entry.detail)

        fill = pass_fill if entry.result == "PASS" else (
            fail_fill if entry.result == "FAIL" else warn_fill)
        result_cell.fill = fill

    _auto_width(ws)
    path = output_dir / "AUDIT_LOG.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_approvals(data: ScheduleData, output_dir: Path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Approvals Required"

    headers = ["Action Type", "Subject", "Detail", "Approval Authority", "Status", "Timestamp"]
    for col, h in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=h)
    _style_header(ws, 1, len(headers))

    for row, req in enumerate(data.approval_requests, 2):
        ws.cell(row=row, column=1, value=req.action_type)
        ws.cell(row=row, column=2, value=req.subject)
        ws.cell(row=row, column=3, value=req.detail)
        ws.cell(row=row, column=4, value=req.approval_authority)
        ws.cell(row=row, column=5, value=req.status)
        ws.cell(row=row, column=6, value=req.timestamp)

    _auto_width(ws)
    path = output_dir / "APPROVALS_REQUIRED.xlsx"
    wb.save(path)
    print(f"  Generated: {path}")


def generate_final_report(data: ScheduleData, reports: list[str], output_dir: Path):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading("OLSM Schedule Build — Final Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("Executive Summary", level=1)

        total_students = len(data.students)
        complete = sum(1 for s in data.students.values() if s.is_schedule_complete())
        conflicted = sum(1 for s in data.students.values() if s.conflicts)
        total_sections = sum(1 for s in data.sections.values() if s.period is not None)
        total_teachers = len(data.teachers)
        g9 = data.students_by_grade_gender(9, Gender.GIRLS)
        b10 = data.students_by_grade_gender(10, Gender.BOYS)
        g_rot = sum(1 for s in g9 if s.rotation_section)
        b_rot = sum(1 for s in b10 if s.rotation_section)

        doc.add_paragraph(f"Total Students: {total_students}")
        doc.add_paragraph(f"Complete Schedules: {complete}/{total_students} ({complete/total_students*100:.1f}%)" if total_students else "No students")
        doc.add_paragraph(f"Students with Conflicts: {conflicted}")
        doc.add_paragraph(f"Total Sections Placed: {total_sections}")
        doc.add_paragraph(f"Total Teachers: {total_teachers}")
        doc.add_paragraph(f"Girls Rotation Coverage: {g_rot}/{len(g9)} ({g_rot/len(g9)*100:.1f}%)" if g9 else "No 9th-grade girls")
        doc.add_paragraph(f"Boys Rotation Coverage: {b_rot}/{len(b10)} ({b_rot/len(b10)*100:.1f}%)" if b10 else "No 10th-grade boys")

        doc.add_heading("Audit Summary", level=1)
        passed = sum(1 for e in data.audit_log if e.result == "PASS")
        failed = sum(1 for e in data.audit_log if e.result == "FAIL")
        warnings = sum(1 for e in data.audit_log if e.result == "WARNING")
        doc.add_paragraph(f"Total Checks: {len(data.audit_log)}")
        doc.add_paragraph(f"Passed: {passed}")
        doc.add_paragraph(f"Failed: {failed}")
        doc.add_paragraph(f"Warnings: {warnings}")

        if data.approval_requests:
            doc.add_heading("Approvals", level=1)
            for req in data.approval_requests:
                doc.add_paragraph(f"[{req.status}] {req.action_type}: {req.subject} — {req.detail}")

        doc.add_heading("Checkpoint Reports", level=1)
        for report in reports:
            doc.add_paragraph(report, style="Normal")
            doc.add_paragraph("")

        path = output_dir / "FINAL_REPORT.docx"
        doc.save(path)
        print(f"  Generated: {path}")

    except ImportError:
        path = output_dir / "FINAL_REPORT.txt"
        with open(path, "w") as f:
            f.write("OLSM Schedule Build — Final Report\n")
            f.write("=" * 50 + "\n\n")
            for report in reports:
                f.write(report + "\n\n")
        print(f"  Generated (txt fallback): {path}")


def generate_all_outputs(data: ScheduleData, reports: list[str],
                         output_dir: str | Path):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"\nGenerating output files to: {output_dir}")
    generate_master_grid(data, output_dir)
    generate_student_schedules(data, output_dir)
    generate_class_rosters(data, output_dir)
    generate_teacher_schedules(data, output_dir)
    generate_audit_log(data, output_dir)
    generate_approvals(data, output_dir)
    generate_final_report(data, reports, output_dir)
    print(f"\nAll 7 output files generated successfully.")
